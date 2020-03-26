#! python3
# customInvitations.py - Generate multiple invitations
# per line on invite list given in argument.

import docx

doc = docx.Document()

# Converts given filename into list of guest strings.
inviteFile = open('guests.txt', 'r')
guests = inviteFile.readlines()

# Create new invitation for each guest.
for guest in guests:

	isLast = '\n' not in guest

	# Remove newline characters.
	if not isLast:
		guest = guest[:-1]

	# Add text and format, as desired.
	doc.add_paragraph("It would be the pleasure to have the company of")
	doc.paragraphs[-1].style = 'Heading 2'

	doc.add_paragraph(guest)
	doc.paragraphs[-1].style = 'Heading 1'

	doc.add_paragraph("at 11010 Memory Lane on the Evening of")
	doc.paragraphs[-1].style = 'Heading 2'

	doc.add_paragraph("April 1st")
	doc.paragraphs[-1].style = 'Heading 3'

	doc.add_paragraph("at 7 o'clock")
	doc.paragraphs[-1].style = 'Heading 4'

	if not isLast:
		doc.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
	
	print('Invitation created for %s.' % guest)

# Save document and post confirmation.
invitationsName = 'invitations.docx'
doc.save(invitationsName)

print('Custom invitations successfully created and saved as %s.' % invitationsName)

